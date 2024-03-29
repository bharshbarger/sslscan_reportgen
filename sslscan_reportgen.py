#!/usr/bin/env python3
"""script to parse output from sslscan and find common issues, then dump into a docx"""
try:
    import argparse
    import docx
    import os
    import re
    import sys
    import time
    from docx.shared import Pt
    from docx.shared import Inches
    from docx.shared import RGBColor    
    from libnmap.parser import NmapParser

except Exception as e:
    print('\n[!] Import(s) failed!: {}'.format(e))

class sslscan_beautifier():
    """a class is probably unnecessary for this script"""
    def __init__(self, args):
        #defaults
        #pass in args. this is messy
        self.args = args
        #start time
        self.start_time = time.time()
        #verbosity explicitly off
        self.verbose = False
        #version
        self.version = 'beta.09_15_2017'
        #file to read
        self.ssl_file = None
        #misconfiguration dictionaries
        self.rc4_dict = {}
        self.sslv2_dict = {}
        self.sslv3_dict = {}
        self.des_dict = {}
        self.tls10_dict = {}
        self.weakbits_dict = {}
        self.heartbleed_dict = {}
        self.md5_dict = {}
        #store sslscan host results here for parsing key, val is ip, results
        self.result_dictionary = {}
        #dump reports here
        self.report_dir = './sslscan_reports/'
        #check for report directory, make it if not found
        if not os.path.exists(self.report_dir):
            os.makedirs(self.report_dir)
        #show user script has started
        print('\n**** SSLScan ReportGen ****\n')

    def check_args(self, parser):
        """Validates supplied args"""
        if self.args.verbose is True: print(\
            '[i] Version: {}\n[i] Options: {}'.format(self.version, parser.parse_args()))
        #make sure -f is supplied
        if not self.args.file:
            print ('No file supplied')
            parser.print_help()
            sys.exit(0)
        #make sure -c is set for file naming purposes
        if not self.args.client:
            print ('Please provide a client name')
            parser.print_help()
            sys.exit(0)
        #clean up client name, remove spaces and stuff
        self.args.client = ''.join(e for e in self.args.client if e.isalnum())

    def read_nmap(self):
        nmap_report = NmapParser.parse_fromfile()
        list443 = []
        list1 = [ a.address for a in nmap_report.hosts if (a.get_open_ports()) and 443 in [b[0] for b in a.get_open_ports()] ]
        for x in list1:
            list443.append(x+":443")
        list8443 = []
        list2 = [ a.address for a in nmap_report.hosts if (a.get_open_ports()) and 8443 in [b[0] for b in a.get_open_ports()] ]
        for x in list2:
            list8443.append(x+":8443")
        mergedlist = list443 + list8443
        mergedlist.sort()
        new1 = '\n'.join(mergedlist)
        return new1


    def run_sslscan(self):
        '''multiprocess here'''
        a = True


    def run_nmap(self):
        #nmap -sV --script ssl-enum-ciphers -p 443 <host>
        a  = True


    def read_file(self):
        """open targets file, splits on ip and puts ip and the result into a dictionary"""
        #strip out ansi control chars (sometimes these appear in colorized output)
        #https://stackoverflow.com/questions/14693701/how-can-i-remove-the-ansi-escape-sequences-from-a-string-in-python
        ansi_escape = re.compile(r'(\x9B|\x1B\[)[0-?]*[ -\/]*[@-~]')
        try:
            with open(self.args.file) as f:
                #use a regex to split the file into sections, delimited by the word Testing
                #also use a regex to strip control chars out as the file is iterated
                for i, result in enumerate(re.findall('Testing(.*?)Testing', ansi_escape.sub('', f.read()), re.S)):
                    #look for the first line from the Testing
                    # BUG for some reason the regex removes 'testing' from the result
                    if 'SSL server' in result:
                        #set up a dictionary key of ip, val of the result
                        self.result_dictionary\
                        [re.findall(r'(?:\d{1,3}\.)+(?:\d{1,3})', result)[0]] = result
        except Exception as e:
            print('\n[!] Couldn\'t open file: {}'.format(e))

    def parse_dict(self):
        """reads thru result_dictionary and search for misconfiguraitons in the scan results"""
        #run through the dictonary containing IPs and the scan results as keys
        for key, val in self.result_dictionary.items():
            #for each line of the scan output
            #print(key)
            for line in val.splitlines():
                #look for RC4
                if 'RC4' in line:
                    self.rc4_dict[key] = line
                if 'SSLv3' in line:
                    self.sslv3_dict[key] = line
                if 'SSLv2' in line:
                    self.sslv2_dict[key] = line
                if 'MD5' in line:
                    self.md5_dict[key] = line
                if 'DES' in line:
                    self.des_dict[key] = line
                if 'TLSv1.0' in line:
                    self.tls10_dict[key] = line
                if 'bits' in line:
                    bits = line.split()[2]
                    if int(bits) < 128:
                        self.weakbits_dict[key] = line
                if 'to heartbleed' in line:
                    if not 'not' in line:
                        self.heartbleed_dict[key] = line

    def gen_report(self):

        print ('\n**** Generating Report: ****')
        print('**** {}sslscan_{}.docx ****\n'.format(self.report_dir, self.args.client))

        """create a document -- to use a template put the docx in the (), like (my_template.docx)"""
        document = docx.Document()


        #header
        heading = document.add_heading(level=3)
        run_heading = heading.add_run('Hosts With Weak Transport Security')
        font = run_heading.font
        font.name = 'Arial'
        font.size = Pt(24)
        font.color.rgb = RGBColor(0xe9, 0x58, 0x23)

        #create output table
        rc4_table = document.add_table(rows=1, cols=2)
        rc4_table.style = 'Table Grid'
        if self.rc4_dict:
            heading_cells = rc4_table.rows[0].cells
            heading_cells[0].text = 'RC4 Hosts'
            heading_cells[1].text = 'Supported Suites'
            for i, r in self.rc4_dict.items():
                cells = rc4_table.add_row().cells
                cells[0].text = str(i)
                cells[1].text = str(r)

        paragraph = document.add_paragraph()
        #run_paragraph = paragraph.add_run('\n')

        sslv2_table =  document.add_table(rows=1, cols=2)
        sslv2_table.style = 'Table Grid'
        if self.sslv2_dict:
            heading_cells = sslv2_table.rows[0].cells
            heading_cells[0].text = 'SSLv2 Hosts'
            heading_cells[1].text = 'Supported Suites'
            for i, r in self.sslv2_dict.items():
                cells = sslv2_table.add_row().cells
                cells[0].text = str(i)
                cells[1].text = str(r)

        paragraph = document.add_paragraph()
        #run_paragraph = paragraph.add_run('\n')

        sslv3_table =  document.add_table(rows=1, cols=2)
        sslv3_table.style = 'Table Grid'
        if self.sslv3_dict:
            heading_cells = sslv3_table.rows[0].cells
            heading_cells[0].text = 'SSLv3 Hosts'
            heading_cells[1].text = 'Supported Suites'
            for i, r in self.sslv3_dict.items():
                cells = sslv3_table.add_row().cells
                cells[0].text = str(i)
                cells[1].text = str(r)

        paragraph = document.add_paragraph()
        #run_paragraph = paragraph.add_run('\n')

        des_table =  document.add_table(rows=1, cols=2)
        des_table.style = 'Table Grid'
        if self.des_dict:
            heading_cells = des_table.rows[0].cells
            heading_cells[0].text = 'DES Hosts'
            heading_cells[1].text = 'Supported Suites'
            for i, r in self.des_dict.items():
                cells = des_table.add_row().cells
                cells[0].text = str(i)
                cells[1].text = str(r)

        paragraph = document.add_paragraph()
        #run_paragraph = paragraph.add_run('\n')

        tls10_table =  document.add_table(rows=1, cols=2)
        tls10_table.style = 'Table Grid'
        if self.tls10_dict:
            heading_cells = tls10_table.rows[0].cells
            heading_cells[0].text = 'TLS v1.0 Hosts'
            heading_cells[1].text = 'Supported Suites'
            for i, r in self.tls10_dict.items():
                cells = tls10_table.add_row().cells
                cells[0].text = str(i)
                cells[1].text = str(r)

        paragraph = document.add_paragraph()
        #run_paragraph = paragraph.add_run('\n')

        weakbits_table =  document.add_table(rows=1, cols=2)
        weakbits_table.style = 'Table Grid'
        if self.weakbits_dict:
            heading_cells = weakbits_table.rows[0].cells
            heading_cells[0].text = 'Weak Key Length Hosts'
            heading_cells[1].text = 'Supported Suites'
            for i, r in self.weakbits_dict.items():
                cells = weakbits_table.add_row().cells
                cells[0].text = str(i)
                cells[1].text = str(r)

        paragraph = document.add_paragraph()
        #run_paragraph = paragraph.add_run('\n')


        heartbleed_table =  document.add_table(rows=1, cols=2)
        heartbleed_table.style = 'Table Grid'
        if self.heartbleed_dict:
            heading_cells = heartbleed_table.rows[0].cells
            heading_cells[0].text = 'Heartbleed Hosts'
            heading_cells[1].text = 'Supported Suites'
            for i, r in self.heartbleed_dict.items():
                cells = heartbleed_able.add_row().cells
                cells[0].text = str(i)
                cells[1].text = str(r)

        paragraph = document.add_paragraph()
        #run_paragraph = paragraph.add_run('\n')


        md5_table =  document.add_table(rows=1, cols=2)
        md5_table.style = 'Table Grid'
        if self.md5_dict:
            heading_cells = md5_table.rows[0].cells
            heading_cells[0].text = 'MD5 Hosts'
            heading_cells[1].text = 'Supported Suites'
            for i, r in self.md5_dict.items():
                cells = md5_table.add_row().cells
                cells[0].text = str(i)
                cells[1].text = str(r)

        document.save('{}sslscan_{}.docx'.format(self.report_dir, self.args.client.strip().encode()))


    def print_summary(self):
        """displays what was found to terminal"""
        if self.rc4_dict:
            print('\n********RC4 Hosts********')
            for k, v in sorted(self.rc4_dict.items()):
                print('{} {}'.format(k, v))
        if self.sslv2_dict:
            print('\n********SSLv2 Hosts********')
            for k, v in sorted(self.sslv2_dict.items()):
                print('{} {}'.format(k, v))
        if self.sslv3_dict:
            print('\n********SSLv3 Hosts********')
            for k, v in sorted(self.sslv3_dict.items()):
                print('{} {}'.format(k, v))
        if self.des_dict:
            print('\n********DES Hosts********')
            for k, v in sorted(self.des_dict.items()):
                print('{} {}'.format(k, v))
        if self.tls10_dict:
            print('\n********TLS v1.0 Hosts********')
            for k, v in sorted(self.tls10_dict.items()):
                print('{} {}'.format(k, v))
        if self.weakbits_dict:
            print('\n********Weak Key Size********')
            for k, v in sorted(self.weakbits_dict.items()):
                print('{} {}'.format(k, v))
        if self.heartbleed_dict:
            print('\n********Heartbleed Hosts********')
            for k, v in sorted(self.heartbleed_dict.items()):
                print('{} {}'.format(k, v))
        if self.md5_dict:
            print('\n********MD5 Hosts********')
            for k, v in sorted(self.md5_dict.items()):
                print('{} {}'.format(k, v))

    def end(self):
        """ending stuff, right now just shows how long script took to run"""
        print('\nCompleted in {:.2f} seconds\n'.format(time.time() - self.start_time))

def main():
    #gather options
    parser = argparse.ArgumentParser()
    parser.add_argument('-c', '--client', metavar='<client>', help='client name')
    parser.add_argument('-f', '--file', metavar='<file>', help='file to read')
    parser.add_argument('-v', '--verbose', help='Optionally enable verbosity', action='store_true')
    parser.add_argument('-x',  '--xml', help='nmap xml output to read')
    args = parser.parse_args()
    run = sslscan_beautifier(args)
    run.check_args(parser)
    run.read_file()
    run.parse_dict()
    run.print_summary()
    run.gen_report()
    run.end()

if __name__ == '__main__':
    main()
